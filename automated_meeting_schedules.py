from office365.sharepoint.client_context import ClientContext
import pandas as pd
from docx import Document
import requests
import numpy as np
import datetime
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx_form import DocxForm
import os
from dateutil import parser as dateparser
import platform

if platform.system() == "Windows":
    from docx2pdf import convert

import re
from datetime import datetime
from tqdm import tqdm
import json
import subprocess
import shutil

import random

script_dir = os.path.dirname(os.path.abspath(__file__))
config_path = os.path.join(script_dir, "automated_meeting_schedules_preferences.json")
with open(config_path) as f:
    config = json.load(f)

base_meetings_folder = config["base_meetings_folder"]
master_spreadsheet_url = config["master_spreadsheet_url"]
meeting_schedule_template_url = config["meeting_schedule_template_url"]



def download_with_progress(url, filename, desc="Downloading"):
    """
    Download a file from a SharePoint edit/view URL using requests, with a progress bar.
    Handles SharePoint redirect to the actual file download link.
    """
    session = requests.Session()
    # Try to get the file with streaming
    response = session.get(url, stream=True, allow_redirects=True)
    # If SharePoint, look for a downloadUrl in the redirected page
    if "text/html" in response.headers.get("Content-Type", ""):
        # Try to extract the download link from the HTML
        match = re.search(r'"downloadUrl":"([^"]+)"', response.text)
        if match:
            download_url = match.group(1).replace('\\u0026', '&')
            response = session.get(download_url, stream=True)
        else:
            raise Exception("Could not find a direct download link in the SharePoint page.")
    total = response.headers.get('content-length')
    total = int(total) if total is not None else None
    with open(filename, "wb") as file, tqdm(
        desc=desc,
        total=total,
        unit='B',
        unit_scale=True,
        unit_divisor=1024,
        disable=(total is None)
    ) as bar:
        for data in response.iter_content(chunk_size=1024):
            size = file.write(data)
            if total is not None:
                bar.update(size)
    return response


# --- Prompt for NCSA if configured ---
try:
    if config.get("include_NCSA_prompt") and random.randint(1, 1) == 1: # 1% chance to show NCSA prompt
        response = input("What is the better NCSA, NBB or Hawk? ").strip()
        if (response.strip().upper() != "NBB") and ("blue beret" not in response.lower()):
            print("Wrong answer. Exiting.")
            exit()
except KeyError:
    ...

# --- Download master spreadsheet from SharePoint ---
master_spreadsheet_path = os.path.join(script_dir, "master_spreadsheet.xlsx")

if config.get("run_in_offline_mode"):
    print("Running in offline mode. Skipping SharePoint download.")
    # Check if the master spreadsheet exists in script_dir
    if not os.path.isfile(master_spreadsheet_path):
        print(f"master_spreadsheet.xlsx not found. Please ensure it is located in {script_dir}.")
        exit()
else:
    print("Running in online mode. Downloading master spreadsheet from SharePoint...")
    master_response = download_with_progress(master_spreadsheet_url, master_spreadsheet_path, desc="Downloading master spreadsheet")
    if master_response.status_code == 200:
        print("Spreadsheet downloaded successfully!")
    else:
        print(f"Failed to download spreadsheet. Check the link permissions. Status code: {master_response.status_code}")
        exit()


# --- Meeting schedule file selection logic ---
meeting_schedule_filename = None

if config.get("run_in_offline_mode"):
    print("Running in offline mode. Skipping SharePoint download.")
    # Search for a file containing "Meeting Schedule" in the name
    for fname in os.listdir():
        if "Meeting Schedule" in fname and fname.lower().endswith(".docx"):
            response = input(f"Use {fname} as the meeting schedule template? [Y/N]: ").strip().upper()
            if response == "Y" or response == "YES":
                meeting_schedule_filename = fname
                break
    # If not found, prompt the user for the file name
    if not meeting_schedule_filename:
        meeting_schedule_filename = input("Enter the full file name for the meeting schedule (including .docx): ").strip()
        if not os.path.isfile(meeting_schedule_filename):
            print(f"File '{meeting_schedule_filename}' not found. Exiting.")
            exit()
else:
    # Online mode: download if not found locally
    for fname in os.listdir():
        if "Meeting Schedule" in fname and fname.lower().endswith(".docx"):
            response = input(f"Use {fname} as the meeting schedule template? [Y/N]: ").strip().upper()
            if response == "Y" or response == "YES":
                meeting_schedule_filename = fname
                break
            else:  # Couldn't find a suitable file, so we will download it
                meeting_schedule_response = download_with_progress(meeting_schedule_template_url, "Meeting Schedule.docx", desc="Downloading meeting schedule template")
                if meeting_schedule_response.status_code == 200:
                    print("Meeting schedule template downloaded successfully!")
                    meeting_schedule_filename = "Meeting Schedule.docx"
                else:
                    meeting_schedule_filename = input("Enter the full file name for the meeting schedule (including .docx): ").strip()
                    if not os.path.isfile(meeting_schedule_filename):
                        print(f"File '{meeting_schedule_filename}' not found. Exiting.")
                        exit()
    # If not found, prompt the user for the file name and/or download
    if not meeting_schedule_filename:
        meeting_schedule_response = download_with_progress(meeting_schedule_template_url, "Meeting Schedule.docx", desc="Downloading meeting schedule template")
        if meeting_schedule_response.status_code == 200:
            print("Meeting schedule template downloaded successfully!")
            meeting_schedule_filename = "Meeting Schedule.docx"
        else:
            meeting_schedule_filename = input("Enter the full file name for the meeting schedule template (including .docx): ").strip()
            if not os.path.isfile(meeting_schedule_filename):
                print(f"File '{meeting_schedule_filename}' not found. Exiting.")
                exit()
doc = Document(os.path.join(script_dir, meeting_schedule_filename))


# Check if the sheet "Current" exists before reading (using pandas)
with pd.ExcelFile(master_spreadsheet_path) as xls:
    # Find a sheet name that matches 'current' (case-insensitive, allow for 'curent' typo)
    sheet_name = None
    for s in xls.sheet_names:
        if s.strip().lower() in ["current"]:
            sheet_name = s
            break
    if not sheet_name:
        print("⚠️ Error: The Excel file does not contain a sheet named 'Current'. Please check your spreadsheet.")
        exit()
    # Now you can safely read the sheet
    df_preview = pd.read_excel(xls, sheet_name=sheet_name, header=None)


# Find the row index where "Meeting" appears
header_row_idx = None
for i, row in df_preview.iterrows():
    if "Meeting" in row.values:
        header_row_idx = i
        break

if header_row_idx is None:
    raise ValueError("Could not find 'Meeting' header in the sheet.")

# Reload the DataFrame using the found header row
df = pd.read_excel(master_spreadsheet_path, sheet_name="Current", header=header_row_idx)
df.columns = df.columns.str.strip()
#print("Columns found in Excel:", list(df.columns))

date_match = re.search(r'(\d{1,2})-(\d{1,2})-(\d{2,4})', meeting_schedule_filename)
if date_match:
    month, day, year = date_match.groups()
    # Pad month and day with zeros if needed
    month = month.zfill(2)
    day = day.zfill(2)
    # Use 4-digit year
    if len(year) == 2:
        year = "20" + year
    target_date = datetime.strptime(f"{month}-{day}-{year}", "%m-%d-%Y")
else:
    # Prompt user for date
    while True:
        user_date = input("Enter the meeting date (MM-DD-YYYY): ").strip()
        # Split and fill in missing parts
        parts = user_date.split('-')
        now = datetime.now()
        try:
            if len(parts) == 3:
                month, day, year = parts
            elif len(parts) == 2:
                month, day = parts
                year = str(now.year)
            elif len(parts) == 1 and parts[0]:
                month = parts[0]
                day = "01"
                year = str(now.year)
            else:
                raise ValueError
            month = month.zfill(2)
            day = day.zfill(2)
            if len(year) == 2:
                year = "20" + year
            target_date = datetime.strptime(f"{month}-{day}-{year}", "%m-%d-%Y")


            break
        except Exception:
            print("Invalid date format. Please enter as MM-DD-YYYY.")

        


found_meeting = False  # Track if a matching meeting is found

for idx, row in df.iterrows():
    if pd.isnull(row["Meeting"]):
        continue
    meeting_date = row["Meeting"]
    # Compare as datetime or parse as string
    match = False
    if isinstance(meeting_date, datetime) and meeting_date.date() == target_date.date():
        match = True
    else:
        try:
            parsed = dateparser.parse(str(meeting_date), fuzzy=True)
            match = parsed.date() == target_date.date()
        except Exception:
            match = False
    if match:
        found_meeting = True
        details = []

        # --- 1. Add details from the same row (to the right of the date cell) ---
        detail = {
            "Class": row.get("Class", ""),
            "Topic": row.get("Topic", ""),
            "Time": row.get("Time", ""),
            "Instructor": row.get("Instructor", "")
        }
        # Only add if at least one field is not empty
        if any(str(v).strip() for v in detail.values()):
            details.append(detail)

        # --- 2. Add details from rows below until a blank row is found ---
        next_idx = idx + 1
        while next_idx < len(df):
            next_row = df.iloc[next_idx]
            # Stop if the row is completely blank
            if next_row.isnull().all():
                break
            detail = {
                "Class": next_row.get("Class", ""),
                "Topic": next_row.get("Topic", ""),
                "Time": next_row.get("Time", ""),
                "Instructor": next_row.get("Instructor", "")
            }
            if any(str(v).strip() for v in detail.values()):
                details.append(detail)
            next_idx += 1

        print(f"Processing meeting on {meeting_date} with details:")
        for detail in details:
            print(f"Class: {detail.get('Class', '')}, "
                  f"Topic: {detail.get('Topic', '')}, "
                  f"Time: {detail.get('Time', '')}, "
                  f"Instructor: {detail.get('Instructor', '')}")


        # --- Remove existing tables before adding a new one ---
        # This will remove all tables in the document
        tables = doc.tables
        for table_obj in tables:
            tbl_element = table_obj._element
            tbl_element.getparent().remove(tbl_element)

        # Create a table with 1 header row and 4 columns
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Agenda', 'Instruction Summary', 'Time', 'Instructor']
        for i, text in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(text)
            run.font.size = Pt(16)
            run.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if hdr_cells[i].paragraphs[0].text != text:
                hdr_cells[i].paragraphs[0].clear()
            # Set background shading for header cell
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'D7D7D7')
            tcPr.append(shd)

        # Set the width of the "Time" column (index 2) to about 1.1 inches 
        for row in table.rows:
            row.cells[1].width = Inches(2.31)
            row.cells[2].width = Inches(1.1)

        for detail in details:
            row_cells = table.add_row().cells
            row_cells[0].text = str(detail.get('Class', ''))
            row_cells[1].text = str(detail.get('Topic', ''))
            row_cells[2].text = str(detail.get('Time', ''))
            row_cells[3].text = str(detail.get('Instructor', ''))

            # Center all text in the row and set font size
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(13)

                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.find(qn('w:tcMar'))
                if tcMar is None:
                    tcMar = OxmlElement('w:tcMar')
                    tcPr.append(tcMar)
                for dir in ('top', 'bottom'):
                    mar = tcMar.find(qn(f'w:{dir}'))
                    if mar is None:
                        mar = OxmlElement(f'w:{dir}')
                        tcMar.append(mar)
                    mar.set(qn('w:w'), '50')  # 0.5pt = 10 twentieths of a point
                    mar.set(qn('w:type'), 'dxa')

        # Also set width for new rows
        for row in table.rows:
            row.cells[2].width = Inches(1.1)

        # Replace the placeholder date in the document
        target_date_str = target_date.strftime("%B %d, %Y")  # e.g., "May 19, 2025"

        target = "[date], 2025"
        for para in doc.paragraphs:
            if target in para.text.lower():
                # Find the start and end of the match in the original text
                start = para.text.lower().index(target)
                end = start + len(target)
                prefix = para.text[:start]
                suffix = para.text[end:]
                para.clear()
                run1 = para.add_run(prefix)
                run1.font.size = Pt(16)
                run1.bold = False
                run2 = para.add_run(target_date_str)
                run2.font.size = Pt(16)
                run2.bold = True
                if suffix:
                    run3 = para.add_run(suffix)
                    run3.font.size = Pt(16)
                    run3.bold = False

        # --- Uniform logic for drill test sign-up phrase ---
        # Find the value in the cell directly below the meeting date, in the "Meeting" column
        uniform_value = ""
        if idx + 1 < len(df):
            uniform_value = str(df.iloc[idx + 1]["Meeting"]).strip().upper()

        drill_phrase = config["drill_test_sign_up_phrase"].strip()

        print(f"Uniform found: {uniform_value}")

        if config["add_drill_test_signup_text_to_abu_uniform_meetings"]:

            # Helper: find paragraph index with the phrase
            def find_paragraph_with_phrase(doc, phrase):
                for i, para in enumerate(doc.paragraphs):
                    if phrase in para.text:
                        return i
                return None

            para_idx = find_paragraph_with_phrase(doc, drill_phrase)

            if "abus" in uniform_value.lower():
                print(f"Adding drill test sign-up phrase. (Edit or disable this in the config file if needed.)")
                # If not present, add it (size 16, bold) before the table
                if para_idx is None:
                    # Insert before the table by adding to the end, then moving it up
                    para = doc.add_paragraph()
                    run = para.add_run(drill_phrase)
                    run.font.size = Pt(16)
                    run.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the paragraph
                    # Move the new paragraph to just before the table
                    doc._body._element.remove(para._element)
                    table_idx = len(doc.paragraphs)  # Table will be added next
                    doc._body._element.insert(table_idx, para._element)
                # If present, do nothing (keep it)
            else:
                # If present, remove it
                if para_idx is not None:
                    p = doc.paragraphs[para_idx]
                    p.clear()
                

        # Determine dropdown_choice based on uniform_value
        dropdown_choice = None
        if uniform_value.lower()  == "pt":
            dropdown_choice = "PT"
        elif uniform_value.lower()  == "blues" or uniform_value.lower()  == "blues with tie":
            dropdown_choice = "Blues with Tie"
        elif "abus" in uniform_value.lower() and "(sleeves rolled)" in uniform_value.lower() or "sleeves rolled" in uniform_value.lower():
            dropdown_choice = "ABUs (sleeves rolled)"
        elif "abus" in uniform_value.lower() and "(sleeves down)" in uniform_value.lower() or "sleeves down" in uniform_value.lower():
            dropdown_choice = "ABUs (sleeves down)"
        elif uniform_value.lower()  == "abu":
            sleeves = input("ABU sleeves rolled or down not specified. Should ABU sleeves be rolled? [Y/N]: ").strip().upper()
            if sleeves == "Y":
                dropdown_choice = "ABUs (sleeves rolled)"
            else:
                dropdown_choice = "ABUs (sleeves down)"


        
        # If dropdown_choice is set, replace [UNIFORM] in the document with this value
        if dropdown_choice:
            for para in doc.paragraphs:
                if "[UNIFORM]" in para.text:
                    prefix, suffix = para.text.split("[UNIFORM]", 1)
                    para.clear()
                    # Add the prefix (font size 16, not bold)
                    run1 = para.add_run(prefix)
                    run1.font.size = Pt(16)
                    run1.bold = False
                    # Add the uniform value (font size 16, bold)
                    run2 = para.add_run(dropdown_choice)
                    run2.font.size = Pt(16)
                    run2.bold = True
                    # Add any suffix (font size 16, not bold)
                    if suffix:
                        run3 = para.add_run(suffix)
                        run3.font.size = Pt(16)
                        run3.bold = False
        
        break  # Stop after processing the first matching meeting

if not found_meeting:
    print(f"No meeting found in the spreadsheet for {target_date.strftime('%m-%d-%Y')}. Exiting.")
    exit()


save_file_date = target_date.strftime("%m-%d-%Y")  


# Delete the master_spreadsheet.xlsx file after processing
if os.path.isfile(master_spreadsheet_path) and not config.get("run_in_offline_mode"):
    try:
        os.remove(master_spreadsheet_path)
    except Exception as e:
        print(f"Could not delete {master_spreadsheet_path}: {e}. Please do not make any changes to it as the script re-downloads it every time it runs.")

    try:
        os.remove(os.path.join(script_dir, "Meeting Schedule.docx"))
    except Exception as e:
        ...



response = input(f"File will be saved as {save_file_date} Meeting Schedule. Continue? [Y/N]: ").strip().upper()
if response != "Y" and response != "YES":
    print("Operation cancelled by user. Exiting.")
    exit()


# Check if the current working directory's name matches save_file_date (is a meeting folder)
current_dir_name = os.path.basename(os.getcwd())
if current_dir_name == save_file_date:
    print(f"I seem to be running in a meeting folder. I will save the document here.")
    save_path = os.getcwd()
else: 
    print(f"Saving the document in the corrisponding folder specified in preferences.")

    base_meetings_folder = config["base_meetings_folder"]
    save_path = os.path.join(base_meetings_folder, save_file_date)
    if not os.path.exists(save_path):
        print(f"Unable to find a Meeting Schedules folder at {save_path}.")
        create_folder = input(f"Create a new folder with this name ({save_file_date})? (otherwise the files will be saved locally to current directory) [Y/N]: ").strip().upper()
        if create_folder == "Y":
            os.makedirs(save_path)
            print(f"Created folder: {save_path}")
        else:
            print("Saving to the current directory instead.")
            save_path = os.getcwd()
            exit()


# Save the modified document
modified_doc_file = f"{save_file_date} Meeting Schedule.docx"
docx_full_path = os.path.join(save_path, modified_doc_file)
doc.save(docx_full_path)
print(f"Document saved as {modified_doc_file} in {save_path}.")


def try_export_pdf(docx_path, save_path, pdf_file_name):
    pdf_path = os.path.join(save_path, pdf_file_name)
    docx2pdf_error = None

    # Only try docx2pdf on Windows
    if platform.system() == "Windows":
        try:
            convert(docx_path, pdf_path)
            print(f"PDF exported as {pdf_file_name} to {pdf_path} using docx2pdf.")
            return True
        except Exception as e:
            docx2pdf_error = e  # Store the error, don't print yet

    # Try LibreOffice (works on Linux, Windows, Mac if installed)
    soffice_path = shutil.which("soffice")
    if soffice_path:
        try:
            subprocess.run([
                soffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_path),
                docx_path
            ], check=True)
            print(f"PDF exported as {pdf_file_name} to {pdf_path} using LibreOffice.")
            return True
        except Exception as e:
            print("PDF export failed. Please ensure LibreOffice is installed.")
            if docx2pdf_error:
                print(f"docx2pdf failed: {docx2pdf_error}")
            print(f"LibreOffice PDF export failed: {e}")
    else:
        if docx2pdf_error:
            print(f"docx2pdf failed: {docx2pdf_error}")

    print("PDF export failed. Please ensure LibreOffice is installed.")
    return False

# Export as PDF (try docx2pdf, then LibreOffice, else skip)

pdf_file_name = f"{save_file_date} Meeting Schedule.pdf"
try_export_pdf(docx_full_path, save_path, pdf_file_name)
